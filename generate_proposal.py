from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 페이지 여백 설정 ──
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ── 스타일 헬퍼 ──
def set_font(run, size=11, bold=False, color=None):
    run.font.name = '맑은 고딕'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    rPr.insert(0, rFonts)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_font(run, 14, bold=True, color=(0, 70, 127))
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
    elif level == 2:
        set_font(run, 12, bold=True, color=(0, 100, 160))
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 3:
        set_font(run, 11, bold=True, color=(40, 40, 40))
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, 10.5)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(16)
    return p

def add_bullet(doc, text, level=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 * level)
    run = p.add_run(text)
    set_font(run, 10.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def shade_cell(cell, hex_color='D9E1F2'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size, bold=bold, color=color)

def add_table(doc, headers, rows, col_widths=None, header_color='1F4E79'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, header_color)
        set_cell_text(cell, h, bold=True, size=10,
                      align=WD_ALIGN_PARAGRAPH.CENTER,
                      color=(255, 255, 255))

    # 데이터
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 0:
                shade_cell(cell, 'EBF3FB')
            align = WD_ALIGN_PARAGRAPH.RIGHT if str(val).replace(',','').replace('-','').replace('+','').replace('%','').replace('.','').isdigit() else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cell, str(val), size=10, align=align)

    # 열 너비
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table

def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════
# 표지
# ══════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run('연 구 과 제 제 안 서')
set_font(run, 22, bold=True, color=(0, 70, 127))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LLM 기반 실시간 생산 의사결정 지원 시스템 개발\n및 스마트 팩토리 적용 연구')
set_font(run, 16, bold=True, color=(30, 30, 30))

doc.add_paragraph()
doc.add_paragraph()

cover_info = [
    ('지원 기관', '한국연구재단 (NRF)'),
    ('사업 유형', '중견연구자지원사업'),
    ('연구 기간', '3년 (1차년도 ~ 3차년도)'),
    ('총 연구비', '900,000천원 (300,000천원/년)'),
    ('연구 분야', '스마트제조 / 인공지능 융합'),
]
for label, value in cover_info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f'{label}: ')
    set_font(r1, 12, bold=True, color=(0, 70, 127))
    r2 = p.add_run(value)
    set_font(r2, 12)

doc.add_page_break()

# ══════════════════════════════════════════
# 목차
# ══════════════════════════════════════════
add_heading(doc, '목  차', 1)
toc_items = [
    ('1.', '연구 개요 및 필요성'),
    ('2.', '연구의 독창성 및 혁신성'),
    ('3.', '세부 연구목표 (Specific Aims)'),
    ('4.', '예산 편성 내역'),
    ('5.', '연구 추진 일정표'),
    ('6.', '기대 성과 및 활용 방안'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    r1 = p.add_run(f'{num}  {title}')
    set_font(r1, 11)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ══════════════════════════════════════════
# 1. 연구 개요 및 필요성
# ══════════════════════════════════════════
add_heading(doc, '1. 연구 개요 및 필요성 (Significance)', 1)
add_divider(doc)

add_heading(doc, '1.1 문제 정의', 2)
add_body(doc,
    '국내 제조업 평균 OEE(Overall Equipment Effectiveness)는 55~65% 수준으로, '
    '세계 수준(85%)과 20~30%p 격차가 존재한다. 이 격차의 핵심 원인은 생산 현장의 '
    '의사결정 지연과 비정형 정보(설비 이력, 작업자 경험, 품질 보고서)가 운영 판단에 '
    '통합되지 못하는 구조적 문제에 있다. 기존 MES/ERP는 정형 데이터만 처리하며, '
    '실시간 이상 상황에서의 대응 판단은 여전히 숙련 작업자의 암묵지에 의존한다.')

add_heading(doc, '1.2 연구 공백 (Research Gap)', 2)
add_body(doc,
    '현재 LLM의 제조 적용 연구는 설계 문서 요약, 코드 생성 등 비운영 영역에 집중되어 있으며, '
    '실시간 생산 루프 내 의사결정 지원에 LLM을 통합한 연구는 극히 드물다. '
    '특히 한국 제조 환경(중소·중견 제조업, 혼류 생산, 고숙련 인력 감소)에 특화된 모델은 전무하다.')

add_heading(doc, '1.3 예비 연구 근거', 2)
add_body(doc,
    '사출성형 라인(INJ-001) 8시간 실측 데이터 분석 결과, 가동률 85.4% / 성능률 84.6% / '
    '품질률 98.0% → OEE 70.8%로 세계 수준 대비 14.2%p 격차를 확인하였다. '
    '성능률 손실의 주원인이 금형 교체 후 안정화 지연(비정형 노하우 의존)임을 실증적으로 식별하였다.')

add_table(doc,
    ['지표', '현재 수준', '세계 수준', '격차'],
    [
        ['가동률 (Availability)', '85.4%', '90%↑', '-4.6%p'],
        ['성능률 (Performance)',  '84.6%', '95%↑', '-10.4%p'],
        ['품질률 (Quality)',      '98.0%', '99%↑', '-1.0%p'],
        ['OEE (종합)',            '70.8%', '85%↑', '-14.2%p'],
    ],
    col_widths=[6, 3.5, 3.5, 3.5]
)

# ══════════════════════════════════════════
# 2. 독창성 및 혁신성
# ══════════════════════════════════════════
add_heading(doc, '2. 연구의 독창성 및 혁신성 (Innovation)', 1)
add_divider(doc)

add_table(doc,
    ['구분', '기존 접근', '본 연구'],
    [
        ['데이터 처리', '정형 센서 데이터', '정형 + 비정형(로그, 매뉴얼, 작업일지) 통합'],
        ['의사결정',   '규칙 기반 알람',    'LLM 추론 기반 맥락적 판단 제안'],
        ['인터페이스', '대시보드 조회',     '자연어 질의응답 (작업자 대화형)'],
        ['지식 보존',  '개인 숙련에 의존',  '암묵지 자동 구조화·축적'],
    ],
    col_widths=[3.5, 5.5, 7.5]
)

# ══════════════════════════════════════════
# 3. Specific Aims
# ══════════════════════════════════════════
add_heading(doc, '3. 세부 연구목표 (Specific Aims)', 1)
add_divider(doc)

aims = [
    (
        'Aim 1 (기초, 1차년도): 제조 도메인 특화 LLM 파인튜닝 및 RAG 파이프라인 구축',
        '협력 기업 3곳에서 5년치 비정형 데이터 수집 및 전처리 후 RAG(Retrieval-Augmented Generation) '
        '아키텍처로 실시간 문서 검색을 통합한다. 제조 도메인 벤치마크 데이터셋을 구축·공개한다.',
        '도메인 QA 정확도 85% 이상, 일반 LLM 대비 20%p 향상'
    ),
    (
        'Aim 2 (메커니즘, 2차년도): 실시간 생산 이상 감지 및 원인 추론 시스템 개발',
        '시계열 이상 감지 모델(Transformer 기반)과 LLM 추론 엔진을 통합하여 고장 원인을 '
        '"5-Why" 구조로 자동 생성하는 Chain-of-Thought 프롬프팅을 설계한다. '
        'SPC 위반 시 자연어 보고서를 자동 생성하는 파이프라인을 구현한다.',
        '이상 감지 후 원인 추론 응답시간 < 3초, 현장 전문가 판단 일치율 80% 이상'
    ),
    (
        'Aim 3 (응용, 3차년도): 작업자 대화형 의사결정 지원 인터페이스 현장 실증',
        '태블릿/HMI 연동 자연어 인터페이스를 개발하고 A/B 테스트(시스템 사용 라인 vs. 미사용 라인) '
        '6개월 비교를 통해 OEE 향상 효과를 실증한다. 작업자 암묵지 자동 구조화 및 지식베이스 축적을 검증한다.',
        'OEE 5%p 이상 향상, 설비 다운타임 15% 감소, 작업자 만족도 4.0/5.0 이상'
    ),
]

for title, desc, kpi in aims:
    add_heading(doc, title, 3)
    add_body(doc, f'[연구 내용] {desc}')
    add_body(doc, f'[성과 지표] {kpi}')

add_heading(doc, '3.1 위험 관리 (Risk Mitigation)', 2)
add_table(doc,
    ['위험 요소', '가능성', '대응 전략 (Plan B)'],
    [
        ['기업 데이터 확보 난항', '중', 'MOU 선제 체결, 공개 데이터셋(PHM, MIMII) 병행 활용'],
        ['LLM 추론 지연 > 3초',  '중', '엣지 배포용 경량화 모델(LoRA) 병행 개발'],
        ['현장 작업자 수용성 저조', '저', 'UX 전문가 참여, 파일럿 라인 단계적 확산'],
        ['비정형 데이터 품질 불균일', '고', '데이터 품질 자동 필터링 파이프라인 1차년도 선행 구축'],
    ],
    col_widths=[5, 2, 9.5]
)

doc.add_page_break()

# ══════════════════════════════════════════
# 4. 예산 편성
# ══════════════════════════════════════════
add_heading(doc, '4. 예산 편성 내역', 1)
add_divider(doc)

add_heading(doc, '4.1 총괄표 (단위: 천원)', 2)
add_table(doc,
    ['비목', '1차년도', '2차년도', '3차년도', '합계', '비율'],
    [
        ['인건비',       '150,000', '150,000', '150,000', '450,000', '50.0%'],
        ['연구재료비',   '55,000',  '40,000',  '30,000',  '125,000', '13.9%'],
        ['연구기자재비', '40,000',  '10,000',  '5,000',   '55,000',  '6.1%'],
        ['연구활동비',   '25,000',  '30,000',  '35,000',  '90,000',  '10.0%'],
        ['위탁연구비',   '15,000',  '20,000',  '15,000',  '50,000',  '5.6%'],
        ['간접비',       '15,000',  '15,000',  '15,000',  '45,000',  '5.0%'],
        ['국제공동연구비', '0',     '35,000',  '50,000',  '85,000',  '9.4%'],
        ['합계',         '300,000', '300,000', '300,000', '900,000', '100%'],
    ],
    col_widths=[4, 2.8, 2.8, 2.8, 2.8, 2]
)

add_heading(doc, '4.2 인건비 세부 (연 150,000천원 × 3년)', 2)
add_table(doc,
    ['구분', '직급', '참여율', '연간 금액(천원)', '역할'],
    [
        ['연구책임자',   '교수',       '30%',  '30,000', '총괄 및 연구 방향 설정'],
        ['공동연구원 A', '교수(AI)',   '20%',  '18,000', '제조 LLM 개발 총괄'],
        ['공동연구원 B', '교수(산업)', '20%',  '18,000', 'OEE·현장 검증 담당'],
        ['전임연구원',   '박사후연구원', '100%', '42,000', 'RAG 파이프라인 전담'],
        ['연구보조원',   '박사과정 2인', '50%', '28,000', '모델 학습·실험'],
        ['연구보조원',   '석사과정 1인', '50%', '14,000', '데이터 수집·전처리'],
    ],
    col_widths=[3, 3.5, 2.2, 3.8, 5]
)

add_heading(doc, '4.3 연구재료비 세부 (합계 125,000천원)', 2)
add_table(doc,
    ['항목', '1차년도', '2차년도', '3차년도', '용도'],
    [
        ['GPU 클라우드 (AWS/Azure)', '25,000', '20,000', '10,000', 'LLM 파인튜닝, 대규모 실험'],
        ['데이터 수집·라벨링',       '15,000', '8,000',  '5,000',  '협력사 데이터 정제·어노테이션'],
        ['LLM API 사용료',           '8,000',  '5,000',  '5,000',  '베이스라인 비교·프롬프트 실험'],
        ['소프트웨어 라이선스',       '5,000',  '5,000',  '5,000',  '개발 툴, DB, 분석 플랫폼'],
        ['소모성 연구재료',           '2,000',  '2,000',  '5,000',  'IIoT 센서 소모품, 배선 자재'],
    ],
    col_widths=[4.5, 2.5, 2.5, 2.5, 5.5]
)

doc.add_page_break()

# ══════════════════════════════════════════
# 5. 연구 추진 일정표
# ══════════════════════════════════════════
add_heading(doc, '5. 연구 추진 일정표', 1)
add_divider(doc)

add_heading(doc, '5.1 전체 Gantt Chart', 2)

gantt_rows = [
    ['[기반] 문헌조사 및 연구 설계',       '■', '■', '',  '',  '',  '',  '',  '',  '',  '',  '',  '' ],
    ['[기반] 협력기업 MOU 체결',            '■', '◆', '',  '',  '',  '',  '',  '',  '',  '',  '',  '' ],
    ['[기반] GPU 서버 구축',                '■', '■', '',  '',  '',  '',  '',  '',  '',  '',  '',  '' ],
    ['[Aim 1] 제조 데이터 수집·전처리',     '',  '■', '■', '',  '',  '',  '',  '',  '',  '',  '',  '' ],
    ['[Aim 1] 도메인 벤치마크 구축',        '',  '',  '■', '■', '',  '',  '',  '',  '',  '',  '',  '' ],
    ['[Aim 1] LLM 파인튜닝',               '',  '',  '■', '■', '',  '',  '',  '',  '',  '',  '',  '' ],
    ['[Aim 1] RAG 파이프라인 구축',         '',  '',  '',  '■', '■', '',  '',  '',  '',  '',  '',  '' ],
    ['[Aim 2] 시계열 이상 감지 모델 개발',  '',  '',  '',  '',  '■', '■', '',  '',  '',  '',  '',  '' ],
    ['[Aim 2] LLM 추론 엔진 통합',          '',  '',  '',  '',  '',  '■', '■', '',  '',  '',  '',  '' ],
    ['[Aim 2] SPC 자동 보고서 파이프라인',  '',  '',  '',  '',  '',  '',  '■', '■', '',  '',  '',  '' ],
    ['[Aim 2] 국제공동연구 (Fraunhofer)',   '',  '',  '',  '',  '',  '◆', '■', '■', '',  '',  '',  '' ],
    ['[Aim 3] 현장 HMI 인터페이스 개발',   '',  '',  '',  '',  '',  '',  '',  '■', '■', '',  '',  '' ],
    ['[Aim 3] 파일럿 라인 배포',            '',  '',  '',  '',  '',  '',  '',  '',  '◆', '■', '',  '' ],
    ['[Aim 3] A/B 테스트 실증 (6개월)',     '',  '',  '',  '',  '',  '',  '',  '',  '',  '■', '■', '' ],
    ['[Aim 3] 최종 성과 분석 및 확산',      '',  '',  '',  '',  '',  '',  '',  '',  '',  '',  '■', '■'],
    ['[공통] 특허 출원',                    '',  '',  '',  '◆', '',  '',  '',  '◆', '',  '',  '',  '◆'],
    ['[공통] SCI 논문 투고',                '',  '',  '',  '◆', '',  '',  '◆', '◆', '',  '',  '◆', '◆'],
    ['[공통] 연차 보고서 제출',             '',  '',  '',  '◆', '',  '',  '',  '◆', '',  '',  '',  '◆'],
]

headers_gantt = ['연구 내용', 'Y1Q1','Y1Q2','Y1Q3','Y1Q4','Y2Q1','Y2Q2','Y2Q3','Y2Q4','Y3Q1','Y3Q2','Y3Q3','Y3Q4']
table = doc.add_table(rows=1 + len(gantt_rows), cols=len(headers_gantt))
table.style = 'Table Grid'

hdr = table.rows[0]
for i, h in enumerate(headers_gantt):
    shade_cell(hdr.cells[i], '1F4E79')
    set_cell_text(hdr.cells[i], h, bold=True, size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color=(255, 255, 255))

for r_idx, row in enumerate(gantt_rows):
    for c_idx, val in enumerate(row):
        cell = table.rows[r_idx + 1].cells[c_idx]
        if r_idx % 2 == 0:
            shade_cell(cell, 'EBF3FB')
        align = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        color = (0, 112, 192) if val == '■' else (192, 0, 0) if val == '◆' else None
        bold = val in ('■', '◆', '▲')
        set_cell_text(cell, val, bold=bold, size=9, align=align, color=color)

# 열 너비
widths = [5.5] + [0.85] * 12
for i, w in enumerate(widths):
    for row in table.rows:
        row.cells[i].width = Cm(w)

doc.add_paragraph()

add_heading(doc, '5.2 차년도별 핵심 마일스톤', 2)
add_table(doc,
    ['분기', '핵심 활동', '주요 산출물'],
    [
        ['Y1Q1', '문헌조사, GPU 서버 발주, MOU 협의',                   '연구계획서 확정'],
        ['Y1Q2', 'MOU 체결 (협력사 2곳), 데이터 수집 착수',             'MOU 협약서, 수집 프로토콜'],
        ['Y1Q3', '비정형 데이터 전처리, LLM 파인튜닝 1차',              '전처리 파이프라인, 학습 데이터셋'],
        ['Y1Q4', 'RAG 완성, 벤치마크 평가, 특허 출원',                  '특허 1건, SCI 1편 투고, 연차보고서'],
        ['Y2Q1', '시계열 Transformer 이상 감지 모델 개발',              '이상 감지 모듈 v1'],
        ['Y2Q2', 'LLM 추론 엔진 통합, Fraunhofer 공동연구 착수',        'CoT 추론 엔진'],
        ['Y2Q3', 'SPC 자동 보고서 생성 완성, 중간 검증',                'SPC 자동화 시스템, SCI 1편 투고'],
        ['Y2Q4', '통합 시스템 검증, 특허 출원',                         '특허 1건, SCI 1편, 연차보고서'],
        ['Y3Q1', 'HMI·태블릿 인터페이스 개발, 파일럿 배포',            '현장 배포 시스템 v1.0'],
        ['Y3Q2', 'A/B 테스트 착수, 암묵지 수집',                        '중간 OEE 비교 리포트'],
        ['Y3Q3', 'A/B 테스트 완료, 공개 워크숍 개최',                   'SCI 1편 투고, 워크숍 자료'],
        ['Y3Q4', '기술이전 협의, 최종 보고서',                          '특허 1건, 최종보고서, 기술이전'],
    ],
    col_widths=[2, 7, 7.5]
)

doc.add_page_break()

# ══════════════════════════════════════════
# 6. 기대 성과 및 활용 방안
# ══════════════════════════════════════════
add_heading(doc, '6. 기대 성과 및 활용 방안', 1)
add_divider(doc)

add_heading(doc, '6.1 학술적 성과', 2)
add_table(doc,
    ['구분', '목표', '세부 내용'],
    [
        ['SCI(E) 논문',    '5편 이상', '제1저자 3편 이상 포함'],
        ['국내 학술지',    '3편 이상', '대한산업공학회지, 한국정밀공학회지'],
        ['국제학회 발표',  '5회 이상', 'NeurIPS, ICML, IISE Annual Conference'],
        ['국내 특허 출원', '3건',      'LLM 추론 엔진, RAG 아키텍처, HMI 인터페이스'],
        ['소프트웨어 등록', '2건',     '이상 감지 모듈, SPC 자동 보고서 시스템'],
    ],
    col_widths=[4, 3, 9.5]
)

add_heading(doc, '6.2 기술 성숙도(TRL) 목표', 2)
add_table(doc,
    ['구성 요소', '연구 전', '1차년도', '2차년도', '3차년도'],
    [
        ['제조 특화 LLM',    'TRL 1', 'TRL 3', 'TRL 4', 'TRL 5'],
        ['RAG 파이프라인',   'TRL 1', 'TRL 4', 'TRL 5', 'TRL 6'],
        ['이상 감지 모듈',   'TRL 2', 'TRL 3', 'TRL 5', 'TRL 6'],
        ['현장 HMI 시스템',  'TRL 1', 'TRL 2', 'TRL 4', 'TRL 6'],
    ],
    col_widths=[5, 2.5, 2.5, 2.5, 2.5]
)

add_heading(doc, '6.3 경제·사회적 파급효과', 2)
add_table(doc,
    ['항목', '산정 근거', '추정 효과'],
    [
        ['OEE 향상 → 생산성 증가', '국내 제조업 OEE 60→65%, 부가가치 기준', '연 약 1.2조원 (중소제조업 5% 도입 시)'],
        ['설비 수리비 절감',        '다운타임 15% 감소, 비계획 유지보수 기준',  '업체당 연 평균 8,000만원'],
        ['불량률 감소',             '품질률 1%p 향상 기준',                    '업체당 연 평균 3,000만원'],
    ],
    col_widths=[4.5, 6, 6]
)

add_heading(doc, '6.4 활용 방안', 2)
add_body(doc, '▶ 단기 (연구 종료 후 1~2년): 협력 제조업체 2곳 직접 기술이전 적용. 핵심 모듈 오픈소스 공개 (GitHub).')
add_body(doc, '▶ 중기 (3~5년): 개발 시스템 패키지화 후 스타트업·SI 기업 이전. KEIT 스마트공장 보급사업 연계. KS 표준화 기여.')
add_body(doc, '▶ 장기 (5년↑): 제조 특화 파운데이션 모델(K-Manufacturing LLM) 초석으로 활용. 범부처 대형과제 기반 기술 연계.')

add_heading(doc, '6.5 인력 양성 계획', 2)
add_table(doc,
    ['구분', '인원', '양성 내용', '기대 진로'],
    [
        ['박사과정',   '2인',    '제조 AI 풀스택 (데이터·모델·시스템)', '대학 교원, 연구소, 제조 AI 스타트업'],
        ['석사과정',   '1인',    'IIoT 데이터 파이프라인, MLOps',      '산업체 데이터 엔지니어'],
        ['박사후연구원', '1인',  '국제공동연구 경험, 논문 주저자',      '해외 포닥 또는 국내 교원'],
        ['학부 인턴',  '연 2인', '제조 AI 기초 실습',                  '대학원 진학 또는 취업'],
    ],
    col_widths=[3.5, 2, 6, 5]
)

# ══════════════════════════════════════════
# 저장
# ══════════════════════════════════════════
output_path = r'C:\Users\user\OneDrive\바탕 화면\AICoding\스마트제조_연구과제_제안서.docx'
doc.save(output_path)
print(f'저장 완료: {output_path}')
